from __future__ import annotations

import io
import re
from dataclasses import dataclass

import pypdfium2 as pdfium
import pytesseract
from docx import Document
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader
from pytesseract import TesseractError, TesseractNotFoundError


@dataclass(frozen=True)
class ExtractedDocument:
    name: str
    text: str


@dataclass(frozen=True)
class DocumentProcessingError(Exception):
    message: str
    status_code: int = 400


def extract_document_text(
    *,
    file_name: str,
    data: bytes,
    fallback_text: str,
    enable_ocr: bool,
    ocr_languages: str,
) -> ExtractedDocument:
    name = file_name.strip() or "document.txt"
    extension = _extension_of(name)

    if extension == "pdf":
        text = _extract_pdf_text(data)
        if text and not looks_unreadable_extracted_text(text):
            return ExtractedDocument(name=name, text=text)

        if not enable_ocr:
            raise DocumentProcessingError(
                "This PDF text cannot be read correctly. Enable OCR to read it as an image.",
            )

        return ExtractedDocument(
            name=name,
            text=_extract_scanned_pdf_text(data, ocr_languages),
        )

    if extension == "docx":
        return ExtractedDocument(name=name, text=_extract_docx_text(data))

    if extension in {"jpg", "jpeg", "png"}:
        if not enable_ocr:
            raise DocumentProcessingError(
                "OCR is turned off. Enable OCR to read image files.",
            )

        return ExtractedDocument(
            name=name,
            text=_extract_image_text(data, ocr_languages),
        )

    if fallback_text.strip():
        return ExtractedDocument(
            name=name,
            text=clean_extracted_text(fallback_text),
        )

    return ExtractedDocument(
        name=name,
        text=clean_extracted_text(data.decode("utf-8", errors="replace")),
    )


def _extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as error:
        raise DocumentProcessingError("Could not read this PDF file.") from error

    return clean_extracted_text("\n\n".join(pages))


def _extract_scanned_pdf_text(data: bytes, ocr_languages: str) -> str:
    try:
        document = pdfium.PdfDocument(data)
        texts: list[str] = []

        for page in document:
            bitmap = page.render(scale=2).to_pil()
            texts.append(_ocr_image(bitmap, ocr_languages))
            page.close()

        document.close()
    except (TesseractError, TesseractNotFoundError):
        raise
    except Exception as error:
        raise DocumentProcessingError(
            "Could not render this scanned PDF for OCR.",
        ) from error

    text = clean_extracted_text("\n\n".join(texts))
    if not text:
        raise DocumentProcessingError(
            "OCR could not find readable text in this scanned PDF.",
        )

    return text


def _extract_docx_text(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as error:
        raise DocumentProcessingError("Could not read this DOCX file.") from error

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    text = clean_extracted_text("\n\n".join(parts))
    if not text:
        raise DocumentProcessingError(
            "This DOCX file does not contain readable text.",
        )

    return text


def _extract_image_text(data: bytes, ocr_languages: str) -> str:
    try:
        image = Image.open(io.BytesIO(data))
    except UnidentifiedImageError as error:
        raise DocumentProcessingError("Could not read this image file.") from error

    text = clean_extracted_text(_ocr_image(image, ocr_languages))
    if not text:
        raise DocumentProcessingError(
            "OCR could not find readable text in this image.",
        )

    return text


def _ocr_image(image: Image.Image, ocr_languages: str) -> str:
    try:
        return pytesseract.image_to_string(image, lang=ocr_languages)
    except TesseractNotFoundError as error:
        raise DocumentProcessingError(
            "Backend OCR requires Tesseract OCR to be installed.",
            500,
        ) from error
    except TesseractError as error:
        raise DocumentProcessingError(
            f"OCR failed. Check that these Tesseract languages are installed: {ocr_languages}.",
            500,
        ) from error


def clean_extracted_text(text: str) -> str:
    cleaned = (
        text.replace("\u00a0", " ")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )
    cleaned = re.sub(r"-\s*\n\s*", "", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    paragraphs = []
    for paragraph in re.split(r"\n{2,}", cleaned):
        normalized = re.sub(r"\n+", " ", paragraph).strip()
        if normalized:
            paragraphs.append(normalized)

    cleaned = "\n\n".join(paragraphs)
    cleaned = re.sub(r"([,;:])(?=\S)", r"\1 ", cleaned)
    cleaned = re.sub(r"([.!?])(?=[A-Za-z0-9\u00c0-\u1ef9])", r"\1 ", cleaned)
    cleaned = re.sub(
        r"([a-z\u00e0-\u1ef9])([A-Z\u00c0-\u1ef8])",
        r"\1 \2",
        cleaned,
    )
    cleaned = re.sub(r"([A-Za-z\u00c0-\u1ef9])(\d)", r"\1 \2", cleaned)
    cleaned = re.sub(r"(\d)([A-Za-z\u00c0-\u1ef9])", r"\1 \2", cleaned)
    cleaned = re.sub(r"(\S)(\[)", r"\1 \2", cleaned)
    cleaned = re.sub(r"(\])(?=\S)", r"\1 ", cleaned)
    cleaned = re.sub(
        r"([a-z]{4,})(and|with|for|from|into|using)(\s+[A-Z])",
        r"\1 \2\3",
        cleaned,
    )
    cleaned = re.sub(r"([a-z]{4,})witha(\s+[A-Z])", r"\1 with a\2", cleaned)

    paragraphs = [_join_common_split_words(item) for item in re.split(r"\n{2,}", cleaned)]
    cleaned = "\n\n".join(paragraphs)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)

    return cleaned.strip()


def looks_unreadable_extracted_text(text: str) -> bool:
    sample = text[:5000]
    visible_characters = sum(1 for char in sample if not char.isspace())

    if visible_characters == 0:
        return False

    suspicious_count = sum(1 for char in sample if _is_suspicious_extracted_char(char))
    suspicious_ratio = suspicious_count / visible_characters
    long_runs = len(
        re.findall(r"[A-Z\u00c0-\u1ef8]{18,}|[A-Za-z\u00c0-\u1ef9]{30,}", sample)
    )
    known_broken_vietnamese = re.search(
        r"H\u20ac|C\u00c6|ngh\u203a|Tr\u00a6|Tu\u00a7|H\u20acN|C\u00e6ng|\u00f7|\u00bc|PH\s+TH|THI\u203aNV",
        sample,
        re.IGNORECASE,
    )

    return bool(
        known_broken_vietnamese
        or suspicious_count >= 12
        or suspicious_ratio >= 0.02
        or (long_runs >= 4 and suspicious_count >= 3)
    )


def _join_common_split_words(text: str) -> str:
    replacements = {
        "tro ng": "trong",
        "tr ong": "trong",
        "kho ng": "khong",
        "kh ong": "khong",
        "n hu": "nhu",
        "nh u": "nhu",
        "nh ung": "nhung",
        "nhu ng": "nhung",
        "c ua": "cua",
        "cu a": "cua",
        "d uoc": "duoc",
        "du oc": "duoc",
        "d en": "den",
        "de n": "den",
        "d e": "de",
        "v oi": "voi",
        "vo i": "voi",
        "v a": "va",
        "c o": "co",
        "n ay": "nay",
        "na y": "nay",
        "p huong": "phuong",
        "ph uong": "phuong",
        "ng hien": "nghien",
        "ngh ien": "nghien",
        "c uu": "cuu",
        "q ua": "qua",
        "qu a": "qua",
        "t he": "the",
        "th e": "the",
        "a nd": "and",
        "an d": "and",
        "w ith": "with",
        "wi th": "with",
        "wit h": "with",
        "f or": "for",
        "fo r": "for",
    }

    fixed = text
    for broken, replacement in replacements.items():
        pattern = re.compile(
            rf"(^|[^A-Za-z]){re.escape(broken)}(?=[^A-Za-z]|$)",
            re.IGNORECASE,
        )
        fixed = pattern.sub(
            lambda match: f"{match.group(1)}{_match_capitalization(match.group(0)[len(match.group(1)):], replacement)}",
            fixed,
        )

    return fixed


def _match_capitalization(original: str, replacement: str) -> str:
    first_letter = next((char for char in original if char.isalpha()), "")
    if not first_letter or not first_letter.isupper():
        return replacement

    return replacement[0].upper() + replacement[1:]


def _is_suspicious_extracted_char(char: str) -> bool:
    suspicious_codepoints = {
        0x00A2,
        0x00A4,
        0x00A5,
        0x00A7,
        0x00A8,
        0x00AC,
        0x00AE,
        0x00AF,
        0x00B1,
        0x00B2,
        0x00B3,
        0x00B5,
        0x00B6,
        0x00B8,
        0x00BC,
        0x00BD,
        0x00BE,
        0x00BF,
        0x00C6,
        0x00D0,
        0x00D7,
        0x00DE,
        0x00E6,
        0x00F0,
        0x00F7,
        0x00FE,
        0x0152,
        0x0153,
        0x0192,
        0x201A,
        0x201E,
        0x2020,
        0x2021,
        0x20AC,
        0xFFFD,
    }
    return ord(char) in suspicious_codepoints


def _extension_of(file_name: str) -> str:
    if "." not in file_name:
        return ""

    return file_name.rsplit(".", 1)[1].lower()
