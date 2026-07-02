from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("Hướng dẫn deploy Text Summarizer lên VPS", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Domain: api.summary.io.vn")
doc.add_paragraph(
    "Tóm tắt đầy đủ các bước từ mua VPS đến khi app chạy được trên Wi‑Fi và 4G."
)
doc.add_paragraph("")

sections = [
    (
        "PHẦN 1 — THUÊ VPS (iNET)",
        [
            "1. Vào iNET → Cloud Server Linux",
            "2. Chọn cấu hình:",
            "   RAM: 2 GB",
            "   CPU: 2 Core",
            "   SSD: 20 GB (40 GB cũng được)",
            "   OS: Ubuntu 22.04-LTS",
            "   Vị trí: Hà Nội hoặc TP.HCM",
            "3. Đăng ký và thanh toán",
            "4. Ghi lại: IP VPS, Port SSH (vd: 24700), User: root, Mật khẩu",
        ],
    ),
    (
        "PHẦN 2 — KẾT NỐI VPS",
        [
            "1. Vào OneDash (panel iNET) → SSH Terminal",
            "2. Kết nối bằng IP VPS và port SSH",
            "3. Bấm Thêm và tiếp tục (lần đầu SSH)",
            "4. Đăng nhập: root + mật khẩu",
            "5. Thấy prompt root@cloud:~# → đã vào server",
        ],
    ),
    (
        "PHẦN 3 — CÀI DOCKER",
        [
            "Trong terminal VPS chạy:",
            "CODE: curl -fsSL https://get.docker.com | sh",
            "CODE: docker --version",
        ],
    ),
    (
        "PHẦN 4 — ĐƯA CODE BACKEND LÊN VPS",
        [
            "1. Trên máy Windows: nén folder backend thành backend.zip",
            "2. OneDash → Quản lý File → /root/ → Tải lên → backend.zip",
            "3. Trên VPS:",
            "CODE: cd /root",
            "CODE: apt install -y unzip",
            "CODE: unzip backend.zip",
            "CODE: ls /root/backend",
            "Phải thấy: docker-compose.yml, Dockerfile, app/, nginx/, requirements.txt",
        ],
    ),
    (
        "PHẦN 5 — TẠO FILE .env",
        [
            "CODE: cd /root/backend",
            "CODE: nano .env",
            "Dán nội dung từ backend/.env trên máy (GEMINI_API_KEY, model...)",
            "Sửa 2 dòng:",
            "CODE: APP_ENV=production",
            "CODE: TRUST_PROXY_HEADERS=true",
            "Lưu: Ctrl+O → Enter → Ctrl+X",
        ],
    ),
    (
        "PHẦN 6 — CHẠY SERVER",
        [
            "CODE: cd /root/backend",
            "CODE: docker compose up -d --build",
            "Chờ 5–10 phút (lần đầu build)",
            "CODE: docker compose ps",
            "CODE: curl http://localhost/health",
            'Kết quả đúng: {"ok":true,"environment":"production"}',
        ],
    ),
    (
        "PHẦN 7 — DNS (TENTEN)",
        [
            "1. Vào TenTen → quản lý domain summary.io.vn",
            "2. Record đang sử dụng → sửa hoặc thêm:",
            "   Tên: api | Loại: A | Giá trị: IP VPS (vd: 103.75.185.222)",
            "3. Lưu → chờ 5–30 phút",
            "4. Test: http://api.summary.io.vn/health",
            '   Thấy {"ok":true,...} → DNS + server OK',
        ],
    ),
    (
        "PHẦN 8 — CẤU HÌNH APP FLUTTER",
        [
            "Tạo/sửa flutter.env trên máy:",
            "CODE: SUMMARY_API_URL=http://api.summary.io.vn",
            "",
            "Chạy app:",
            'CODE: cd "C:\\Users\\HP\\Downloads\\JOB APP\\my_first_app"',
            "CODE: flutter run --dart-define-from-file=flutter.env",
            "",
            "Build APK:",
            "CODE: flutter build apk --release --dart-define-from-file=flutter.env",
        ],
    ),
    (
        "SƠ ĐỒ TỔNG THỂ",
        [
            "User (Wi‑Fi / 4G)",
            "    ↓",
            "api.summary.io.vn (DNS TenTen → IP VPS)",
            "    ↓",
            "VPS: Docker nginx :80 → FastAPI :8787",
            "    ↓",
            "Gemini API (tóm tắt)",
        ],
    ),
    (
        "CHECKLIST NHANH",
        [
            "☐ Thuê VPS (Ubuntu 22.04, 2GB RAM, 2 CPU)",
            "☐ SSH vào VPS",
            "☐ Cài Docker",
            "☐ Upload backend.zip → giải nén",
            "☐ Tạo .env (production)",
            "☐ docker compose up -d --build",
            "☐ curl localhost/health → ok",
            "☐ DNS api → IP VPS",
            "☐ api.summary.io.vn/health → ok",
            "☐ flutter.env → http://api.summary.io.vn",
            "☐ Chạy / build app",
        ],
    ),
    (
        "LƯU Ý QUAN TRỌNG",
        [
            "Không cần mở port modem nhà (server trên VPS)",
            "Không DNS trỏ IP LAN 192.168.x.x (chỉ trỏ IP VPS)",
            "VPS trial 7 ngày: nhớ gia hạn trước khi hết hạn",
            "Sau reboot VPS: cd /root/backend && docker compose up -d",
            "Trước Play Store: làm HTTPS (https://api.summary.io.vn)",
        ],
    ),
    (
        "BƯỚC TIẾP THEO (TÙY CHỌN)",
        [
            "1. HTTPS — Certbot + ssl.conf trên VPS",
            "2. Gia hạn VPS khi hết trial",
            "3. Auth + quota khi có nhiều user trả phí",
        ],
    ),
]

for heading, lines in sections:
    doc.add_heading(heading, level=1)
    for line in lines:
        if line.startswith("CODE: "):
            text = line.replace("CODE: ", "")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif line.startswith("   "):
            doc.add_paragraph(line.strip(), style="List Bullet")
        elif line.startswith("☐"):
            doc.add_paragraph(line, style="List Bullet")
        elif line == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

out = r"C:\Users\HP\Downloads\JOB APP\my_first_app\Huong_dan_deploy_VPS_Text_Summarizer.docx"
doc.save(out)
print(f"Saved: {out}")
